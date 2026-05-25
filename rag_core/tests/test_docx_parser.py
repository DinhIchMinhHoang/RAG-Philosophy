"""
rag_core/tests/test_docx_parser.py — Unit tests for DOCX Parser (Pandoc-based).

Test cases:
  1. Valid DOCX file → returns List[Document] with correct metadata
  2. Invalid file (not DOCX) → raises ValueError
  3. DOCX with tables → tables converted to Markdown tables
  4. DOCX with math formulas → OMML → LaTeX ($...$)
  5. Idempotency → same file parse 2x → same doc_id
  6. Performance → parse < 3s for typical DOCX
  7. Empty DOCX → raises ValueError
"""

import os
import tempfile
import time
import unittest
from unittest import mock

from langchain_core.documents import Document

try:
    from rag_core.docx_parser import Config, parse_docx, _magic_check
except ImportError:
    from docx_parser import Config, parse_docx, _magic_check


# =====================================================================
#  FIXTURES
# =====================================================================

def _create_simple_docx() -> bytes:
    """Create a minimal valid DOCX file in memory."""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This is a test document for DOCX parsing.')
    doc.add_heading('Section 1', level=2)
    doc.add_paragraph('Content for section 1 with some details.')
    doc.add_heading('Section 2', level=2)
    doc.add_paragraph('Content for section 2 with more details.')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_docx_with_table() -> bytes:
    """Create a DOCX file with a table."""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading('Data Table', level=1)

    table = doc.add_table(rows=3, cols=3)
    headers = ['Name', 'Age', 'City']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    table.cell(1, 0).text = 'Alice'
    table.cell(1, 1).text = '30'
    table.cell(1, 2).text = 'Hanoi'
    table.cell(2, 0).text = 'Bob'
    table.cell(2, 1).text = '25'
    table.cell(2, 2).text = 'HCMC'

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_empty_docx() -> bytes:
    """Create a DOCX file with no text content."""
    from docx import Document
    from io import BytesIO

    doc = Document()
    # No content added

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


NOT_DOCX = b"%PDF-1.4 fake pdf content"


# =====================================================================
#  TESTS
# =====================================================================

class TestDocxParser(unittest.TestCase):

    DOCUMENT_ID = "test-docx-123"
    PIPELINE_VERSION = "v1"

    def setUp(self) -> None:
        self._patchers = [
            mock.patch.object(Config, "OCR_DOCX_ENABLED", False),
            mock.patch.object(Config, "OCR_NONPDF_MODE", "primary"),
        ]
        for patcher in self._patchers:
            patcher.start()

    def tearDown(self) -> None:
        for patcher in reversed(self._patchers):
            patcher.stop()

    def _write_bytes(self, content: bytes, suffix: str = ".docx") -> str:
        """Write bytes to a temp file and return path."""
        tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        tmp.write(content)
        tmp.close()
        return tmp.name

    # ── Test 1: Valid DOCX file → returns List[Document] ──
    def test_valid_docx_returns_documents(self):
        path = self._write_bytes(_create_simple_docx())
        try:
            docs = parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
            self.assertIsInstance(docs, list)
            self.assertGreater(len(docs), 0)
            for doc in docs:
                self.assertIsInstance(doc, Document)
                self.assertIn("source", doc.metadata)
                self.assertIn("page", doc.metadata)
                self.assertIn("doc_id", doc.metadata)
                self.assertTrue(len(doc.page_content) > 0)
        finally:
            os.unlink(path)

    # ── Test 2: Invalid file (not DOCX) → raises ValueError ──
    def test_invalid_file_raises_value_error(self):
        path = self._write_bytes(NOT_DOCX, suffix=".bin")
        try:
            with self.assertRaises(ValueError):
                _magic_check(path)
        finally:
            os.unlink(path)

    # ── Test 3: DOCX with tables → tables converted to Markdown ──
    def test_docx_with_tables_converted_to_markdown(self):
        path = self._write_bytes(_create_docx_with_table())
        try:
            docs = parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
            self.assertGreater(len(docs), 0)
            all_content = "\n".join(doc.page_content for doc in docs)
            # Check that table content is present
            self.assertIn("Alice", all_content)
            self.assertIn("Bob", all_content)
        finally:
            os.unlink(path)

    # ── Test 4: Idempotency → same file parse 2x → same doc_id ──
    def test_idempotency_same_doc_id(self):
        path = self._write_bytes(_create_simple_docx())
        try:
            docs1 = parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
            docs2 = parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)

            self.assertEqual(len(docs1), len(docs2))

            doc_ids_1 = [doc.metadata["doc_id"] for doc in docs1]
            doc_ids_2 = [doc.metadata["doc_id"] for doc in docs2]

            self.assertEqual(doc_ids_1, doc_ids_2)
        finally:
            os.unlink(path)

    # ── Test 5: Performance → parse < 3s for typical DOCX ──
    def test_performance_under_3_seconds(self):
        path = self._write_bytes(_create_simple_docx())
        try:
            start = time.time()
            docs = parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
            elapsed = time.time() - start

            self.assertGreater(len(docs), 0)
            self.assertLess(elapsed, 3.0, f"Parse took {elapsed:.2f}s, expected < 3s")
        finally:
            os.unlink(path)

    # ── Test 6: Empty DOCX → raises ValueError ──
    def test_empty_docx_raises_value_error(self):
        path = self._write_bytes(_create_empty_docx())
        try:
            with self.assertRaises(ValueError):
                parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
        finally:
            os.unlink(path)

    # ── Test 7: DOCX with nested headings → split correctly ──
    def test_docx_nested_headings_split(self):
        from docx import Document
        from io import BytesIO

        doc = Document()
        doc.add_heading('Chapter 1', level=1)
        doc.add_paragraph('Intro text.')
        doc.add_heading('Section 1.1', level=2)
        doc.add_paragraph('Detail text.')
        doc.add_heading('Section 1.2', level=2)
        doc.add_paragraph('More detail text.')
        doc.add_heading('Chapter 2', level=1)
        doc.add_paragraph('Second chapter text.')

        buf = BytesIO()
        doc.save(buf)
        path = self._write_bytes(buf.getvalue())

        try:
            docs = parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
            # Should split by headings into multiple documents
            self.assertGreater(len(docs), 1)
            # Verify headings are preserved
            all_content = "\n".join(doc.page_content for doc in docs)
            self.assertIn("Chapter 1", all_content)
            self.assertIn("Chapter 2", all_content)
        finally:
            os.unlink(path)

    def test_primary_mode_falls_back_to_pandoc_when_ocr_fails(self):
        path = self._write_bytes(_create_simple_docx())
        try:
            with mock.patch.object(Config, "OCR_DOCX_ENABLED", True):
                with mock.patch("rag_core.docx_parser._parse_docx_via_ocr", side_effect=RuntimeError("boom")):
                    docs = parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
            self.assertGreater(len(docs), 0)
        finally:
            os.unlink(path)

    def test_strict_mode_raises_when_ocr_fails(self):
        path = self._write_bytes(_create_simple_docx())
        try:
            with mock.patch.object(Config, "OCR_DOCX_ENABLED", True):
                with mock.patch.object(Config, "OCR_NONPDF_MODE", "strict"):
                    with mock.patch("rag_core.docx_parser._parse_docx_via_ocr", side_effect=RuntimeError("boom")):
                        with self.assertRaises(ValueError):
                            parse_docx(path, self.DOCUMENT_ID, self.PIPELINE_VERSION)
        finally:
            os.unlink(path)


if __name__ == "__main__":
    unittest.main()

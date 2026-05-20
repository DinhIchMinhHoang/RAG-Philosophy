"""
docx_parser.py — DOCX Parser for RAG pipeline.

Converts .docx files into List[Document] (same interface as HybridPDFParser.parse_pdf).

Pipeline:
  1. Magic number check (PK\\x03\\x04\\ = ZIP/OOXML)
  2. Load Document (python-docx)
  3. Extract: paragraphs, tables, math, images, comments
  4. Build virtual pages (page break / word count fallback)
  5. OCR images concurrent (OllamaOCREngine — reuse from step1_parser)
  6. Merge OCR results into pages
  7. del doc + gc.collect() (memory cleanup)
  8. Return List[Document] with metadata {source, page, doc_id}
"""

from __future__ import annotations

import base64
import gc
import io
import os
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from langchain_core.documents import Document as LCDocument
from PIL import Image

try:
    from .common.logging_utils import configure_logging, get_logger
    from .config import Config
    from .step1_parser import OllamaOCREngine, MarkdownSanitizer
except ImportError:  # pragma: no cover
    from common.logging_utils import configure_logging, get_logger
    from config import Config
    from step1_parser import OllamaOCREngine, MarkdownSanitizer

configure_logging()
logger = get_logger(__name__)


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@dataclass
class ParsedPage:
    page_idx: int
    content: str
    headings: list[str] = field(default_factory=list)
    image_refs: list[dict] = field(default_factory=list)
    comments: list[str] = field(default_factory=list)
    raw_omml: list[str] = field(default_factory=list)


def parse_docx(file_path: str) -> List[LCDocument]:
    """
    Entry point — Parse DOCX → List[Document].

    Args:
        file_path: Absolute path to .docx file.

    Returns:
        List[LCDocument] with metadata {source, page, doc_id}.
        Same interface as HybridPDFParser.parse_pdf().

    Raises:
        ValueError: If file signature is not ZIP/OOXML (.docx).
    """
    _magic_check(file_path)

    doc = Document(file_path)
    logger.info("[DocxParser] Loaded: %s", file_path)

    pages = _build_virtual_pages(doc)

    try:
        images = _extract_images(doc)
        if images:
            logger.info("[DocxParser] OCR %d images concurrent", len(images))
            ocr_results = _ocr_images_concurrent(images)
            _merge_ocr_into_pages(pages, ocr_results)
    except Exception as exc:
        logger.warning("[DocxParser] Image extraction/OCR failed: %s. Continuing without images.", exc)

    # Skip comments extraction for now (optional feature)
    # comments = _extract_comments(doc)
    # if comments:
    #     _append_comments_to_pages(pages, comments)

    del doc
    gc.collect()

    return _pages_to_documents(pages, os.path.basename(file_path))


def _magic_check(file_path: str) -> None:
    with open(file_path, 'rb') as f:
        header = f.read(4)
    if header[:4] != b'PK\x03\x04':
        raise ValueError(
            "Unsupported file format. Expected .docx (OOXML/ZIP), "
            "but file signature does not match. Please save as .docx first."
        )


def _build_virtual_pages(doc: Document) -> list[ParsedPage]:
    all_content: list[str] = []
    all_headings: list[str] = []

    for para in doc.paragraphs:
        text, heading = _parse_paragraph(para)
        if heading:
            all_headings.append(heading)
        if text:
            all_content.append(text)

        math_text = _parse_math_from_para(para)
        if math_text:
            all_content.append(math_text)

    for table in doc.tables:
        table_md = _serialize_table_md(table)
        if table_md:
            all_content.append(table_md)

    page_breaks = _detect_page_breaks(doc)

    if page_breaks:
        return _split_by_pagebreaks(all_content, all_headings, page_breaks)
    else:
        return _split_by_wordcount(all_content, all_headings, 700)


def _parse_paragraph(para) -> tuple[str, str | None]:
    style = para.style.name or ""

    if "Heading" in style:
        try:
            level = int(style.split()[-1])
        except (ValueError, IndexError):
            level = 2
        return "", f"{'#' * level} {para.text}"

    if para.runs:
        first_run = para.runs[0]
        if first_run.bold and first_run.font.size:
            try:
                if first_run.font.size.pt >= 14:
                    return "", f"## {para.text}"
            except Exception:
                pass

    text = para.text.strip()
    if not text:
        return "", None

    if _is_list_item(para):
        text = f"- {text}"

    text = _extract_hyperlink_text(para, text)

    return text, None


def _is_list_item(para) -> bool:
    return para._element.find(
        f'.//{{{W}}}numPr'
    ) is not None


def _extract_hyperlink_text(para, current_text: str) -> str:
    hyperlinks = para._element.findall(f'.//{{{W}}}hyperlink')
    for hl in hyperlinks:
        link_text = "".join(
            r.text or "" for r in hl.iter(f'{{{W}}}t')
        )
        if link_text and link_text not in current_text:
            current_text = f"{current_text} [link: {link_text}]"
    return current_text


def _serialize_table_md(table) -> str:
    if not table.rows:
        return ""

    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""

    col_count = len(rows[0])
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join(["---"] * col_count) + " |"
    body = "\n".join("| " + " | ".join(row) + " |" for row in rows[1:])

    return f"{header}\n{separator}\n{body}\n"


def _parse_math_from_para(para) -> str:
    omml_elements = para._element.findall(f'.//{{{W}}}omath')
    if not omml_elements:
        return ""
    parts = []
    for elem in omml_elements:
        latex, _ = _parse_math(elem)
        if latex and not latex.startswith("[FORMULA"):
            parts.append(f"${latex}$")
        elif latex.startswith("[FORMULA"):
            parts.append(latex)
    return " ".join(parts)


def _parse_math(element) -> tuple[str, str | None]:
    try:
        from omml2latex import convert_omml
        from lxml import etree
        omml_str = etree.tostring(element, encoding='unicode')
        latex = convert_omml(omml_str)
        return latex, None
    except ImportError:
        logger.warning("[DocxParser] omml2latex not installed. Math formulas will be placeholder.")
        return f"[FORMULA_{uuid.uuid4().hex[:4]}]", None
    except Exception as exc:
        logger.error("[DocxParser] Math conversion failed: %s", exc)
        return f"[FORMULA_{uuid.uuid4().hex[:4]}]", None


def _extract_images(doc: Document) -> list[dict]:
    images = []
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            try:
                img_bytes = rel.target_part.blob
                mime = rel.target_part.content_type
                if mime == 'image/jpeg':
                    img_bytes = _convert_jpg_to_png(img_bytes)
                images.append({"bytes": img_bytes, "mime": mime})
            except Exception as exc:
                logger.warning("[DocxParser] Skip unsupported image (clipboard/ole): %s", exc)
                continue
    return images


def _convert_jpg_to_png(jpg_bytes: bytes) -> bytes:
    img = Image.open(io.BytesIO(jpg_bytes))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _ocr_images_concurrent(images: list[dict]) -> list[str]:
    engine = OllamaOCREngine()
    results = [None] * len(images)

    def ocr_one(idx_img: tuple[int, dict]) -> tuple[int, str]:
        idx, img = idx_img
        b64 = base64.b64encode(img["bytes"]).decode('ascii')
        text = engine.ocr_page_image(b64)
        return idx, text or ""

    max_workers = Config.OLLAMA_MAX_WORKERS
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {
            executor.submit(ocr_one, (i, img)): i
            for i, img in enumerate(images)
        }
        for future in as_completed(futures):
            page_idx = futures[future]
            try:
                idx, text = future.result(timeout=Config.VLM_TIMEOUT_SECONDS)
                results[idx] = text
            except Exception as exc:
                logger.warning("[DocxParser] OCR failed for image %d: %s", page_idx, exc)
                results[page_idx] = ""

    return results


def _detect_page_breaks(doc: Document) -> list[int]:
    page_breaks = []
    for idx, para in enumerate(doc.paragraphs):
        for br in para._element.iter(f'{{{W}}}br'):
            if br.get(f'{{{W}}}type') == 'page':
                page_breaks.append(idx)
                break
    return page_breaks


def _split_by_pagebreaks(contents: list[str], headings: list[str], breaks: list[int]) -> list[ParsedPage]:
    pages = []
    page_idx = 0
    start = 0
    break_points = breaks + [len(contents)]

    for bp in break_points:
        page_content = "\n\n".join(contents[start:bp])
        if page_content.strip():
            pages.append(ParsedPage(
                page_idx=page_idx,
                content=page_content,
                headings=headings,
            ))
            page_idx += 1
        start = bp

    return pages


def _split_by_wordcount(contents: list[str], headings: list[str], words_per_page: int = 700) -> list[ParsedPage]:
    all_text = "\n\n".join(contents)
    words = all_text.split()

    pages = []
    page_idx = 0

    for i in range(0, len(words), words_per_page):
        chunk_words = words[i:i + words_per_page]
        page_content = " ".join(chunk_words)
        if page_content.strip():
            pages.append(ParsedPage(
                page_idx=page_idx,
                content=page_content,
                headings=headings,
            ))
            page_idx += 1

    return pages


def _merge_ocr_into_pages(pages: list[ParsedPage], ocr_results: list[str]) -> None:
    if not ocr_results:
        return
    ocr_text = "\n\n".join(r for r in ocr_results if r)
    if pages:
        pages[0].content = pages[0].content + "\n\n" + ocr_text


def _append_comments_to_pages(pages: list[ParsedPage], comments: list[str]) -> None:
    if not pages or not comments:
        return
    comment_text = "\n\n".join(f"[Comment]: {c}" for c in comments)
    pages[-1].content = pages[-1].content + "\n\n" + comment_text


def _extract_comments(doc: Document) -> list[str]:
    try:
        if not hasattr(doc, 'part') or not hasattr(doc.part, 'comments'):
            return []
        comments_part = doc.part.comments
        if comments_part is None:
            return []
        return [comment.text for comment in comments_part if comment.text]
    except Exception as exc:
        logger.warning("[DocxParser] Failed to extract comments: %s", exc)
        return []


def _pages_to_documents(pages: list[ParsedPage], source: str) -> List[LCDocument]:
    docs = []
    for page in pages:
        clean_content = MarkdownSanitizer.sanitize(page.content)
        if not clean_content:
            continue
        docs.append(LCDocument(
            page_content=clean_content,
            metadata={
                "source": source,
                "page": page.page_idx + 1,
                "doc_id": str(uuid.uuid4()),
            }
        ))
    logger.info("[DocxParser] → %d Document objects", len(docs))
    return docs


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <file.docx>")
        sys.exit(1)

    docs = parse_docx(sys.argv[1])
    print(f"Parsed: {len(docs)} pages")
    for doc in docs:
        print(f"--- Page {doc.metadata['page']} ({len(doc.page_content)} chars)")
        print(doc.page_content[:300])